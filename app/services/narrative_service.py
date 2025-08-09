import os
import logging
from datetime import datetime
import json
from app import db
from app.models import Narrative, Grant, Organization
from app.services.ai_service import ai_service
import docx
from io import BytesIO

logger = logging.getLogger(__name__)

def create_or_update_narrative(grant_id, additional_content=None):
    """
    Create or update a narrative for a grant
    
    Args:
        grant_id (int): The grant ID
        additional_content (str, optional): Additional content for the narrative
        
    Returns:
        Narrative: The created or updated narrative
    """
    try:
        # Get the grant and organization data
        grant = Grant.query.get(grant_id)
        if not grant:
            raise ValueError(f"Grant not found with ID: {grant_id}")
        
        org = Organization.query.first()
        if not org:
            raise ValueError("Organization profile not found")
        
        # Check if a narrative already exists
        narrative = Narrative.query.filter_by(grant_id=grant_id).first()
        
        # Generate the narrative content
        case_for_support = additional_content or org.case_for_support
        # Generate narrative using AI service
        sections = ['need', 'program', 'outcomes', 'budget_rationale']
        narrative_result = ai_service.generate_narrative(
            grant.to_dict(), 
            org.to_dict(), 
            sections
        )
        narrative_content = narrative_result if narrative_result else {}
        
        if narrative:
            # Update existing narrative
            narrative.content = narrative_content
            narrative.last_updated = datetime.now()
        else:
            # Create new narrative
            narrative = Narrative(
                grant_id=grant_id,
                content=narrative_content,
                created_at=datetime.now(),
                last_updated=datetime.now()
            )
            db.session.add(narrative)
        
        db.session.commit()
        return narrative
    
    except Exception as e:
        db.session.rollback()
        logger.error(f"Error creating/updating narrative: {str(e)}")
        raise

def generate_narrative_doc(narrative_id):
    """
    Generate a Word document from a narrative
    
    Args:
        narrative_id (int): The narrative ID
        
    Returns:
        BytesIO: The Word document as a BytesIO object
    """
    try:
        # Get the narrative
        narrative = Narrative.query.get(narrative_id)
        if not narrative:
            raise ValueError(f"Narrative not found with ID: {narrative_id}")
        
        # Get related grant
        grant = Grant.query.get(narrative.grant_id)
        if not grant:
            raise ValueError(f"Grant not found with ID: {narrative.grant_id}")
        
        # Get organization
        org = Organization.query.first()
        if not org:
            raise ValueError("Organization profile not found")
        
        # Create a new Word document
        doc = docx.Document()
        
        # Add title
        doc.add_heading(f"Grant Narrative: {grant.title}", 0)
        
        # Add organization and grant information
        doc.add_heading("Organization Information", 1)
        doc.add_paragraph(f"Organization: {org.name}")
        doc.add_paragraph(f"Mission: {org.mission}")
        
        doc.add_heading("Grant Information", 1)
        doc.add_paragraph(f"Funder: {grant.funder}")
        doc.add_paragraph(f"Amount: ${grant.amount:,.2f}" if grant.amount else "Amount: Not specified")
        doc.add_paragraph(f"Due Date: {grant.due_date.strftime('%B %d, %Y')}" if grant.due_date else "Due Date: Not specified")
        
        # Add narrative content
        doc.add_heading("Narrative", 1)
        
        # Split narrative into paragraphs and add each
        paragraphs = narrative.content.split("\n\n")
        for para in paragraphs:
            if para.strip():
                doc.add_paragraph(para.strip())
        
        # Save the document to a BytesIO object
        docio = BytesIO()
        doc.save(docio)
        docio.seek(0)
        
        return docio
    
    except Exception as e:
        logger.error(f"Error generating narrative document: {str(e)}")
        raise
